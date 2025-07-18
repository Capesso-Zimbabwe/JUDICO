"""Script to create a sample Word document with quotes for testing.

This script reads quotes from sample_quotes.txt and creates a Word document
with one quote per paragraph for testing the import_quotes command.

Requires python-docx package: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Path to the sample quotes text file
TXT_FILE = os.path.join(os.path.dirname(__file__), 'sample_quotes.txt')

# Path for the output Word document
DOCX_FILE = os.path.join(os.path.dirname(__file__), 'sample_quotes.docx')

def create_quotes_docx():
    """Create a Word document with quotes from the text file."""
    # Check if the text file exists
    if not os.path.exists(TXT_FILE):
        print(f"Error: {TXT_FILE} not found.")
        return False
    
    try:
        # Read quotes from the text file
        with open(TXT_FILE, 'r', encoding='utf-8') as f:
            quotes = [line.strip() for line in f if line.strip()]
        
        # Create a new Word document
        doc = Document()
        
        # Add a title
        title = doc.add_heading('JUDICO Daily Quotes', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add an introduction paragraph
        intro = doc.add_paragraph(
            'This document contains quotes to be displayed in the JUDICO application. '
            'Each quote will be shown on a different day throughout the year.'
        )
        intro.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add a section break
        doc.add_paragraph('---')
        
        # Add each quote as a separate paragraph
        for quote in quotes:
            p = doc.add_paragraph(quote)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Add style to the paragraph
            for run in p.runs:
                run.font.size = Pt(11)
        
        # Save the document
        doc.save(DOCX_FILE)
        print(f"Successfully created {DOCX_FILE} with {len(quotes)} quotes.")
        return True
    
    except Exception as e:
        print(f"Error creating Word document: {str(e)}")
        return False

if __name__ == '__main__':
    create_quotes_docx()